import os
from pathlib import Path
from typing import List

from loguru import logger


def parse_file(file_path: str) -> str:
    """Parse a document file and return its text content."""
    path = Path(file_path)
    if not path.exists():
        raise FileNotFoundError("File not found: {}".format(file_path))

    suffix = path.suffix.lower()
    parsers = {
        ".txt": _parse_txt,
        ".md": _parse_txt,
        ".csv": _parse_txt,
        ".json": _parse_txt,
        ".jsonl": _parse_txt,
        ".pdf": _parse_pdf,
        ".docx": _parse_docx,
        ".html": _parse_html,
        ".htm": _parse_html,
    }

    parser = parsers.get(suffix)
    if parser is None:
        logger.warning("Unsupported file type '{}', trying plain text read.".format(suffix))
        parser = _parse_txt

    try:
        text = parser(str(path))
    except Exception as e:
        logger.error("Failed to parse {}: {}".format(file_path, e))
        raise

    return _clean_text(text)


def _parse_txt(file_path: str) -> str:
    encodings = ["utf-8", "gbk", "gb2312", "latin-1"]
    for encoding in encodings:
        try:
            with open(file_path, "r", encoding=encoding) as f:
                return f.read()
        except UnicodeDecodeError:
            continue
    raise UnicodeDecodeError("Unable to decode file with tried encodings: {}".format(file_path))


def _parse_pdf(file_path: str) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as e:
        raise ImportError("pypdf is required for PDF parsing") from e

    reader = PdfReader(file_path)
    parts = []
    for page in reader.pages:
        try:
            text = page.extract_text()
            if text:
                parts.append(text)
        except Exception as e:
            logger.warning("Failed to extract PDF page: {}".format(e))
    return "\n".join(parts)


def _parse_docx(file_path: str) -> str:
    try:
        from docx import Document
    except ImportError as e:
        raise ImportError("python-docx is required for DOCX parsing") from e

    doc = Document(file_path)
    parts = []
    for para in doc.paragraphs:
        if para.text:
            parts.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text for cell in row.cells if cell.text]
            if row_text:
                parts.append(" | ".join(row_text))
    return "\n".join(parts)


def _parse_html(file_path: str) -> str:
    try:
        from trafilatura import extract
    except ImportError as e:
        raise ImportError("trafilatura is required for HTML parsing") from e

    with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
        html = f.read()
    text = extract(html)
    return text or ""


def _clean_text(text: str) -> str:
    if not text:
        return ""
    # Normalize line breaks
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    # Collapse multiple blank lines
    lines = [line.strip() for line in text.split("\n")]
    cleaned = []
    prev_blank = False
    for line in lines:
        blank = line == ""
        if blank and prev_blank:
            continue
        cleaned.append(line)
        prev_blank = blank
    return "\n".join(cleaned).strip()


def list_supported_files(directory: str) -> List[str]:
    """List all supported document files under a directory recursively."""
    supported = {".txt", ".md", ".csv", ".json", ".jsonl", ".pdf", ".docx", ".html", ".htm"}
    files = []
    for root, _, filenames in os.walk(directory):
        for name in filenames:
            if Path(name).suffix.lower() in supported:
                files.append(os.path.join(root, name))
    return sorted(files)
